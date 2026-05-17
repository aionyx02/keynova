"""Generate a generic Prompt Engineering initialization template (.docx).

This script distills the CLAUDE.md + docs/ pattern used in Keynova into a
project-agnostic template that other projects can adopt to bootstrap their
own retrieval-first, ADR-governed documentation system.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "Prompt_Engineering_Init_Template.docx"


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F4F8")
    pPr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


def build_doc() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Prompt Engineering 初始化模板")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = subtitle.add_run("Retrieval-First Docs + CLAUDE.md + ADR Governance — 跨專案通用版")
    s.italic = True
    s.font.size = Pt(12)
    s.font.color.rgb = RGBColor(0x55, 0x60, 0x77)

    doc.add_paragraph()

    add_heading(doc, "0. 使用方式", 1)
    add_paragraph(doc, "目標：讓任何新專案在 30 分鐘內建立可被 AI agent（如 Claude Code）正確讀取、自動同步、具備 ADR 治理的文件骨架。")
    add_numbered(doc, [
        "把本模板列出的目錄樹複製到新專案根目錄。",
        "把所有 [PROJECT_NAME]、[STACK]、[GOAL] 等佔位符替換成實際內容。",
        "保留 frontmatter 的 type / status / context_policy 欄位（agent 用來判斷檢索策略）。",
        "提交第一版後執行一次 docs:refresh（若有腳本）或手動驗證 index.md 對齊。",
        "每個 session 結束前依「Session Close」章節更新 active / current / completed。",
    ])

    add_heading(doc, "1. 核心原則", 1)
    add_bullets(doc, [
        "Retrieval-First：永遠不要一次注入所有文件。先讀 index → memory/current → tasks/active，其餘按 intent 取最小章節。",
        "優先級順序：safety > correctness > rollbackability > testability > performance > speed。",
        "ADR 治理：AI 只能建立狀態為「提議」的 ADR；唯有開發者能標記為「接受」。",
        "文件強制同步：任何任務狀態變更、架構/安全/測試變更，必須在同一 change-set 更新對應文件。",
        "Auto-Archive：完成的 task group（100% [x]）必須在同一 change-set 從 backlog/active 搬到 completed。",
        "衝突解析：以「當前對話的開發者指令 > tasks/active > memory/current > 已接受的 ADR > security > architecture > 歷史 session」為序。",
        "Markdown 是事實來源，但文件本身要小到能 always-retrievable（避免膨脹）。",
    ])

    add_heading(doc, "2. 必要檔案目錄樹", 1)
    add_code(doc, """[REPO_ROOT]/
├── CLAUDE.md                    # session-start bootstrap（必載入）
└── docs/
    ├── CLAUDE.md                # ADR + governance 細則
    ├── index.md                 # 文件路由與 retrieval policy
    ├── project.md               # 穩定的專案事實（stack、平台、目標）
    ├── architecture.md          # 系統架構（按需檢索）
    ├── security.md              # 信任邊界、權限、敏感資料
    ├── testing.md               # 測試策略、覆蓋率、CI 指令
    ├── decisions.md             # ADR 索引（指向 adr/）
    ├── tasks.md                 # tasks/ 索引根
    ├── memory.md                # memory/ 索引根
    ├── tasks/
    │   ├── active.md            # 當下執行中的任務
    │   ├── backlog.md           # 未開工 / 進行中（含 roadmap）
    │   ├── blocked.md           # 阻擋條件、安全 gate
    │   └── completed.md         # 完成歷史
    ├── memory/
    │   ├── current.md           # 當前 working memory（短期）
    │   ├── sessions/            # 每次 session log（按日期）
    │   └── archive/             # 長期歷史（不視為 current 指示）
    └── adr/
        ├── 0000-template.md     # ADR 標準 9 節格式
        └── NNNN-kebab-title.md  # 每個決策一檔
""")

    add_heading(doc, "3. 根目錄 CLAUDE.md 模板", 1)
    add_paragraph(doc, "放在 repo root，會在 agent session 開始時自動載入。保持精簡（< 80 行），細則放 docs/CLAUDE.md。")
    add_code(doc, """---
type: agent_bootstrap
status: active
priority: p0
updated: YYYY-MM-DD
context_policy: always_retrievable
owner: project
---

# CLAUDE.md

> Auto-loaded at session start. Detailed governance and ADR rules are in `docs/CLAUDE.md`.

## Session Start (Retrieval-First)

1. Read `docs/index.md` first (single source for doc routing).
2. Read `docs/memory/current.md` (current strategy and constraints).
3. Read `docs/tasks/active.md` (what to do now).
4. Read `docs/tasks/blocked.md` only when planning, approval, security, or risky actions are involved.
5. Retrieve additional documents by user intent. Never inject all docs into context at once.

## Session Close (Mandatory Sync)

Before final response or commit:

1. Update task state in `docs/tasks/{active,backlog,blocked,completed}.md`.
2. **Auto-archive completed groups (mandatory)**: when a task group reaches 100% `[x]`,
   move the whole section out of `backlog.md` / `active.md` into `completed.md` in the
   same change-set. Do not leave clusters of `[x]` accumulating in `backlog.md`.
3. Update `docs/memory/current.md` (what changed, next step, risks).
4. If architecture/security/testing behavior changed, update the corresponding doc.
5. If decision boundary changed, add/update ADR in `docs/adr/` and refresh `docs/decisions.md`.
6. Run `npm run docs:refresh` (or equivalent) to prevent documentation drift.

## Project Overview

[PROJECT_NAME] is [ONE_LINE_DESCRIPTION].
Primary goal: [GOAL].
Supported platforms / runtimes: [PLATFORMS].

## Git Workflow Rules

```
main  <- always releasable
  <- dev (integration)
      <- feature/<name> (work branches)
```

- Do not merge without explicit user confirmation.
- Create feature branches from `dev`.
- Show diffs and pass checks before merge.

## Build Commands

```bash
[INSTALL_CMD]            # e.g. npm install
[DEV_CMD]                # e.g. npm run dev
[BUILD_CMD]              # e.g. npm run build
[LINT_CMD]               # e.g. npm run lint
[TEST_CMD]               # e.g. npm test
[DOCS_REFRESH_CMD]       # e.g. npm run docs:refresh
```

## Documentation Entry Points

- `docs/index.md` - documentation router and retrieval policy
- `docs/project.md` - stable project facts
- `docs/tasks/active.md` - active work only
- `docs/memory/current.md` - short working memory
- `docs/CLAUDE.md` - ADR and governance rules
""")

    add_heading(doc, "4. docs/CLAUDE.md 模板（治理細則）", 1)
    add_code(doc, """# AI Agent Governance and ADR Policy

Applies to project governance, ADR workflow, and documentation update behavior.

## 1. Core Principle

Safety > correctness > rollbackability > testability > performance > speed.

## 2. Retrieval-First Documentation Rule

- Never inject all docs into prompt at once.
- Always start from `docs/index.md`, `docs/memory/current.md`, `docs/tasks/active.md`.
- Retrieve additional files only by user intent. Use smallest relevant heading section.
- Do not treat `docs/tasks/completed.md` or `docs/memory/archive/*` as current instructions.

## 3. ADR Authority Rule

- AI can create ADR with status `proposed` only.
- AI cannot mark ADR as accepted.
- Implementation that changes architecture/security/data contracts must wait for developer acceptance.

## 4. Documentation Conflict Resolution

When docs conflict, use this order:

1. Explicit developer instruction in current conversation
2. `docs/tasks/active.md`
3. `docs/memory/current.md`
4. Accepted ADR (`docs/adr/*`)
5. `docs/security.md`
6. `docs/architecture.md`
7. Session/archive logs

## 5. Mandatory Documentation Sync

After meaningful changes, update the right file immediately:

- Task status change -> `docs/tasks/*`
- Working strategy/next step -> `docs/memory/current.md`
- Architecture changes -> `docs/architecture.md`
- Security boundary changes -> `docs/security.md`
- Test strategy changes -> `docs/testing.md`
- Decision boundary changes -> `docs/adr/*` + `docs/decisions.md`

### 5a. Auto-Archive Completed Task Groups (mandatory)

Rule: the moment a task group becomes 100% `[x]`, in the same change-set:

1. Remove the entire group section from `backlog.md` (or `active.md`).
2. Append it to `docs/tasks/completed.md` with `(COMPLETE YYYY-MM-DD)` suffix on header.
3. Update Mainline History / Current Execution Order lists.
4. Re-run docs refresh so index and frontmatter stay in sync.

Do not:
- Leave a completed group in `backlog.md` "for later cleanup".
- Move only part of a group; either fully complete (move whole section) or partial (stay).

## 6. Auto-Update Guardrail

Before commit, run the docs refresh script. It should:
1. Update metadata/index automatically.
2. Block commit when code changed but project-state docs were not updated.

## 7. ADR Trigger Checklist

Create ADR before implementation when any apply:

- New/removed core dependency
- Async model / IPC / queue / actor boundary changes
- Security boundary or path/network permissions changes
- Data format / schema / public contract changes
- Major algorithmic / indexing / storage model changes

## 8. Minimal Agent Workflow

1. Classify user intent.
2. Retrieve minimal relevant docs.
3. Implement smallest safe change.
4. Run relevant tests/checks.
5. Update docs via the refresh script.
6. Report with file paths and remaining risks.
""")

    add_heading(doc, "5. 統一 Frontmatter 規格", 1)
    add_paragraph(doc, "所有 docs/*.md 都應該帶 YAML frontmatter，agent 才能依此判斷檢索策略。欄位定義：")
    add_table(
        doc,
        ["欄位", "可選值", "意義"],
        [
            ["type", "agent_bootstrap / agent_policy / docs_index / project_overview / working_memory / memory_index_root / task_index / task_index_root / task_blockers / task_history / architecture_spec / security_policy / testing_policy / decision_index", "agent 用來判斷檔案角色"],
            ["status", "active / backlog / completed / blocked / proposed / accepted", "目前狀態"],
            ["priority", "p0 / p1 / p2 / p3", "重要性"],
            ["updated", "YYYY-MM-DD", "最後修改日期（絕對日期）"],
            ["context_policy", "always_retrievable / retrieve_when_planning / retrieve_when_debugging / retrieve_only / archive", "什麼情境下才注入"],
            ["owner", "project / team-name / individual", "維護負責人"],
            ["tags", "[kebab-case, …]", "選用，輔助檢索"],
        ],
    )

    add_heading(doc, "6. 各檔案內容模板", 1)

    add_heading(doc, "6.1 docs/index.md", 2)
    add_code(doc, """---
type: docs_index
status: active
priority: p0
updated: YYYY-MM-DD
context_policy: always_retrievable
owner: project
---

# [PROJECT_NAME] Documentation Index

## Why this exists

Use this file as the first lookup step. Goal is retrieval-first context, not full-doc prompt injection.

## Default Read Order

1. `docs/index.md`
2. `docs/memory/current.md`
3. `docs/tasks/active.md`
4. `docs/tasks/blocked.md` (only when needed)
5. Additional files by intent

## Retrieval Policy

- Do not read all docs recursively.
- Prefer smallest relevant heading section.
- `completed` and `archive` files are historical context only.
- Conflict order: tasks/active.md + memory/current.md > decisions > sessions/archive.

## Intent Routing

| Intent | Primary docs |
|---|---|
| What should I do now? | memory/current.md, tasks/active.md, tasks/blocked.md |
| Implementation | tasks/active.md, architecture.md, targeted code |
| Security / permission | tasks/blocked.md, security.md, relevant ADR |
| Testing / regression | testing.md, tasks/active.md |
| Historical question | memory/sessions/*, memory/archive/* |

## Context Budget (12k example)

- User request: 1k
- Current memory: 1k
- Active tasks: 2k
- Relevant architecture: 2k
- Relevant code snippets: 4k
- Recent actions/audit: 1k
- Response reserve: 1k

## Document Map

| Path | Type | Status | Context policy | Purpose |
|---|---|---|---|---|
| `docs/CLAUDE.md` | agent_policy | active | always_retrievable | ADR and governance |
| `docs/project.md` | project_overview | active | always_retrievable | Stable project facts |
| `docs/tasks/active.md` | task_index | active | always_retrievable | Current execution tasks |
| `docs/tasks/backlog.md` | task_index | backlog | retrieve_when_planning | Future tasks |
| `docs/tasks/blocked.md` | task_blockers | active | retrieve_when_planning | Blocking constraints |
| `docs/tasks/completed.md` | task_history | completed | archive | Completed history |
| `docs/memory/current.md` | working_memory | active | always_retrievable | Short-term memory |
| `docs/architecture.md` | architecture_spec | active | retrieve_only | System architecture |
| `docs/security.md` | security_policy | active | retrieve_when_planning | Security boundary |
| `docs/testing.md` | testing_policy | active | retrieve_when_debugging | Testing strategy |
| `docs/decisions.md` | decision_index | active | retrieve_only | ADR index |
""")

    add_heading(doc, "6.2 docs/project.md", 2)
    add_code(doc, """---
type: project_overview
status: active
priority: p1
updated: YYYY-MM-DD
context_policy: always_retrievable
owner: project
---

# [PROJECT_NAME] Project Overview

## Product

[ONE_PARAGRAPH 產品定義 + 目標族群 + 核心價值]

## Goals

- [可量化目標 1]
- [可量化目標 2]
- [可量化目標 3]

## Stack

- [Frontend]
- [Backend]
- [Storage]
- [其他關鍵依賴]

## Platform Targets

- [Platform 1 + 最低版本]
- [Platform 2 + 最低版本]

## Engineering Priorities

1. Safety and correctness
2. Rollback and testability
3. Performance and developer speed

## Current Strategy

- [Feature-first / Refactor-first / Safety-first 取一]
- Retrieval-first docs workflow.
- Markdown remains source of truth without full-doc prompt dumping.
""")

    add_heading(doc, "6.3 docs/memory/current.md", 2)
    add_code(doc, """---
type: working_memory
status: active
priority: p0
updated: YYYY-MM-DD
context_policy: always_retrievable
owner: project
---

# Current Project Memory

## Current Strategy

- [短文：當下走哪條路 / 限制 / 取捨]

## Current Focus

- [Phase / Track 名稱 與狀態]

## Important Constraints

- [硬性限制 1，例如「LLM 不得直接執行未授權的 shell」]
- [硬性限制 2]

## Last Confirmed Progress

- YYYY-MM-DD: [一句話交付重點 + 對應檔案路徑]

## Next Step

- [下一個動作；不要寫到下一階段；保持 <= 5 行]

## Known Risks

- [風險與緩解策略]
""")

    add_heading(doc, "6.4 docs/memory.md（索引）", 2)
    add_code(doc, """---
type: memory_index_root
status: active
priority: p0
updated: YYYY-MM-DD
context_policy: always_retrievable
owner: project
---

# Project Memory Index

This file is index-only. Do not store full session history here.

## Memory Sources

- Current working memory: `docs/memory/current.md`
- Session logs: `docs/memory/sessions/`
- Archive: `docs/memory/archive/`

## Rules

- Put only active and short-term facts in `current.md`.
- Move old session logs to `sessions/`.
- Move large historical narratives to `archive/`.
- Do not treat archive as current instruction unless explicitly requested.
""")

    add_heading(doc, "6.5 docs/tasks.md（索引）", 2)
    add_code(doc, """---
type: task_index_root
status: active
priority: p0
updated: YYYY-MM-DD
context_policy: always_retrievable
owner: project
---

# [PROJECT_NAME] Tasks Index

This file is index-only. Do not treat it as the full task source.

Current strategy: [Feature First, Refactor Second / 其他]

## Task Sources

- Active now: `docs/tasks/active.md`
- Backlog: `docs/tasks/backlog.md`
- Blocked / safety constraints: `docs/tasks/blocked.md`
- Completed history: `docs/tasks/completed.md`

## Reading Policy

- For implementation, read `active.md` first.
- For planning and safety, include `blocked.md`.
- `completed.md` is historical only.
- Do not load all task files unless explicitly requested.
""")

    add_heading(doc, "6.6 docs/tasks/active.md", 2)
    add_code(doc, """---
type: task_index
status: active
priority: p0
updated: YYYY-MM-DD
context_policy: always_retrievable
owner: project
tags: [active]
last_change: [一句話：最近交付了什麼]
---

# Active Tasks

## Strategy

[一句話：現階段走哪條路 / 為什麼]

## Current Execution Order

1. [TRACK_CODE].A — [子任務描述]
2. [TRACK_CODE].B — [子任務描述]
3. ...

## Next Phase Proposal (YYYY-MM-DD, awaiting selection)

- [Phase X — Track Y]：[一句話價值說明]

## Recent Execution Notes

- YYYY-MM-DD: [Track 完成摘要：檔案路徑 / 測試數 / 待補]

See `docs/tasks/backlog.md` for detailed breakdown.
See `docs/tasks/blocked.md` for gating constraints.
See `docs/tasks/completed.md` for history.
""")

    add_heading(doc, "6.7 docs/tasks/backlog.md", 2)
    add_code(doc, """---
type: task_index
status: backlog
priority: p1
updated: YYYY-MM-DD
context_policy: retrieve_when_planning
owner: project
tags: [roadmap]
---

# Backlog Tasks

> 已完成的 task groups 已搬到 `docs/tasks/completed.md`，此檔只列尚未動工或進行中的工作。
> 一個 group 達 100% `[x]` 時必須在同次更新立即搬到 `completed.md`。

## Mainline History (see completed.md)

完成依序：
1. ...
2. ...

## Phase Proposal

| Phase | Track | 焦點 | ADR 需求 |
|---|---|---|---|
| N | TRACK.X | [描述] | 否 / ADR-NNN |

## TRACK.X — [Track 名稱]

- [ ] TRACK.X.A — [子任務]
- [ ] TRACK.X.B — [子任務]
""")

    add_heading(doc, "6.8 docs/tasks/blocked.md", 2)
    add_code(doc, """---
type: task_blockers
status: active
priority: p0
updated: YYYY-MM-DD
context_policy: retrieve_when_planning
owner: project
tags: [security, blockers]
---

# Blocked Tasks and Safety Constraints

## [BLOCKED_FEATURE_NAME]

Status: blocked

Reason:
- [為什麼此功能被阻擋]

Blocked by:
- ADR-NNN accepted
- [其他前置條件]

Do not:
- [明確禁止 agent 採取的行動]

Allowed before unblock:
- [可先做的 ADR 草擬 / scaffold]
""")

    add_heading(doc, "6.9 docs/tasks/completed.md", 2)
    add_code(doc, """---
type: task_history
status: completed
priority: p2
updated: YYYY-MM-DD
context_policy: archive
owner: project
---

# Completed Tasks

> 歷史紀錄。Agent 預設不需檢索；只在追溯實作緣由時讀取。

## TRACK.X — [Track 名稱] (COMPLETE YYYY-MM-DD)

- [x] TRACK.X.A — [子任務] — [關鍵交付：檔案路徑 / 測試覆蓋]
- [x] TRACK.X.B — [子任務] — [關鍵交付]
""")

    add_heading(doc, "6.10 docs/architecture.md（大綱）", 2)
    add_code(doc, """---
type: architecture_spec
status: active
priority: p1
updated: YYYY-MM-DD
context_policy: retrieve_only
owner: project
---

# [PROJECT_NAME] System Architecture

> Retrieval policy: 只在實作 / 除錯 / 架構決策相關需求時擷取最小章節。

## 1. 系統概覽
### 1.1 分層架構（ASCII 圖）
### 1.2 技術棧

## 2. 模組邊界
### 2.1 前端
### 2.2 後端
### 2.3 IPC / 通訊
### 2.4 儲存層

## 3. 主要資料流程（含序列圖文字版）

## 4. 並發 / 非同步模型

## 5. 擴充點（plugin / registry）

## 6. 啟動 / 關閉 生命週期

## 7. 已知技術債（連結對應 task track）
""")

    add_heading(doc, "6.11 docs/security.md（大綱）", 2)
    add_code(doc, """---
type: security_policy
status: active
priority: p0
updated: YYYY-MM-DD
context_policy: retrieve_when_planning
owner: project
---

# [PROJECT_NAME] 安全模型

## 1. 安全優先級
Safety > correctness > rollback > test > perf > speed

## 2. 信任邊界（[不信任區域] / [信任邊界] / [信任區域]）

## 3. 檔案路徑安全
- 強制 canonicalize
- 防止路徑穿越
- 不跟隨 symlink 逃逸
- 允許的讀寫路徑清單

## 4. 敏感資料處理
- 禁止記錄清單
- Log 遮蔽規則
- 禁止存放敏感資料的位置

## 5. 網路存取
- 目前允許清單
- TLS / timeout / checksum 規則

## 6. Agent / LLM 工具安全
- ToolPermissionGate
- LLM 輸出不信任原則
- Prompt injection 防護

## 7. IPC / API 安全
- Payload 驗證
- 來源信任

## 8. 修改安全模型時的要求（必須先 ADR）

## 9. 已知安全限制（與計劃）
""")

    add_heading(doc, "6.12 docs/testing.md（大綱）", 2)
    add_code(doc, """---
type: testing_policy
status: active
priority: p1
updated: YYYY-MM-DD
context_policy: retrieve_when_debugging
owner: project
---

# [PROJECT_NAME] 測試策略

## 1. 測試原則
- 驗證行為，不驗證實作細節
- 不為通過測試而改測試
- 不在 production code 加 pub-only-for-test
- 測試不含敏感資料
- 改核心行為 / 資料格式 / 安全邊界 / public API 必須新增或更新測試

## 2. 測試指令
- 後端：[CMD]
- 前端：[CMD]
- E2E：[CMD]

## 3. 測試類型與適用情境（unit / integration / regression / security / performance / migration / manual）

## 4. 覆蓋率目標（模組對覆蓋率表）

## 5. 重要測試案例（按模組列）

## 6. 效能測試基準（含具體量測表）

## 7. 平台限制與已知問題

## 8. 測試資料與 Fixture 規則

## 9. CI 執行清單
""")

    add_heading(doc, "6.13 docs/decisions.md（ADR 索引）", 2)
    add_code(doc, """---
type: decision_index
status: active
priority: p1
updated: YYYY-MM-DD
context_policy: retrieve_only
owner: project
---

# decisions.md — ADR Index

> Retrieval policy: 先由 `docs/index.md` 判斷是否需要 ADR，再精準讀取對應單一 ADR 檔案。

## ADR 流程快速參考

- ADR 格式：`adr/NNNN-kebab-title.md`（9 節標準格式）
- 完整模板：`adr/0000-template.md`
- 治理規則：`docs/CLAUDE.md`
- AI agent 不得自行將 ADR 狀態改為「接受」

### 狀態定義

| 狀態 | 意義 |
|------|------|
| 提議 | 尚未允許實作 |
| 接受 | 開發者已審閱，可開始實作 |
| 拒絕 | 此方案被拒絕 |
| 廢棄 | 曾接受但已不再適用 |
| 取代 | 被新 ADR 取代 |

## ADR 清單

| 編號 | 標題 | 狀態 | 日期 |
|------|------|------|------|
| [0001](adr/0001-xxx.md) | [標題] | 接受 | YYYY-MM-DD |
""")

    add_heading(doc, "6.14 docs/adr/0000-template.md（ADR 標準格式）", 2)
    add_code(doc, """# [決策標題]

**狀態：** 提議
**日期：** YYYY-MM-DD
**決策者：** 開發者 / AI agent / 團隊
**相關文件：**
- docs/architecture.md
- docs/security.md
- docs/testing.md

---

## 1. Context（技術背景）
- 現在的問題是什麼？
- 為什麼現有設計不足？
- 輸入規模 / 未來規模？
- 瓶頸發生在 CPU / memory / disk I/O / network / UI / DX？
- 不處理會有什麼風險？

## 2. Constraints（系統限制與邊界）
- 平台、runtime、記憶體、啟動時間、資料量、安全、網路、背景任務限制

## 3. Alternatives Considered（替代方案分析）
至少列兩個方案。每個方案附：優缺點 / 效能分析（時間 / 空間 / I/O / 啟動 / 維護成本）/ 風險。

## 4. Decision（最終決策）
選擇 / 原因 / 犧牲 / 未來擴充空間 / feature flag / migration / rollback。

## 5. Consequences（系統影響與副作用）
正面、負面/技術債、對使用者、對開發者、對測試、對安全性的影響。

## 6. Implementation Plan（實作計畫）
- 步驟列表（ADR 未接受前不得修改 production code）

## 7. Rollback Plan（回滾策略）
- Feature flag 開關
- 資料格式回滾
- 殘留檔案清理
- 驗證退回是否成功

## 8. Validation Plan（驗證方式）
| 測試類型 | 覆蓋目標 | 指令 / 步驟 |
|---------|---------|-----------|
| Unit | ... | ... |
| Integration | ... | ... |
| Regression | ... | ... |
| Security | ... | ... |
| Performance | ... | ... |
| Manual | ... | ... |

## 9. Open Questions（未解問題）
- [ ] ...
""")

    add_heading(doc, "7. Session Start / Session Close 工作流程", 1)
    add_heading(doc, "7.1 Session Start（agent 啟動時自動執行）", 2)
    add_numbered(doc, [
        "讀 docs/index.md（路由）",
        "讀 docs/memory/current.md（當前策略 + 限制）",
        "讀 docs/tasks/active.md（要做什麼）",
        "依 intent 決定是否要再讀 blocked.md / security.md / architecture.md / 對應 ADR",
        "不要遞迴讀整個 docs/",
    ])
    add_heading(doc, "7.2 Session Close（回覆 / commit 前強制執行）", 2)
    add_numbered(doc, [
        "更新 docs/tasks/{active,backlog,blocked,completed}.md 狀態",
        "若有 task group 100% [x] → 同一 change-set 搬到 completed.md（auto-archive）",
        "更新 docs/memory/current.md（last progress / next step / risks）",
        "若架構 / 安全 / 測試行為變更 → 對應檔同步更新",
        "若決策邊界變更 → 新增 ADR（status=提議）+ 更新 decisions.md",
        "執行 docs:refresh（或等價腳本）阻止文件漂移",
    ])

    add_heading(doc, "8. Auto-Archive 規則（重要）", 1)
    add_paragraph(doc, "task group = backlog.md / active.md 中的 section header，其子項為 [ ] / [x]。例如 `## PERF.1`, `## FEAT.11`, `## AGENT.1`。")
    add_paragraph(doc, "規則：當 group 100% 變成 [x] 的那個瞬間，在同一 change-set：")
    add_numbered(doc, [
        "把整個 group section 從 backlog.md（或 active.md）移除",
        "Append 到 docs/tasks/completed.md，header 加 `(COMPLETE YYYY-MM-DD)` 後綴；保留每個子項的實作備註",
        "更新 backlog.md 的 Mainline History / active.md 的 Current Execution Order",
        "重跑 docs:refresh 讓 index.md 與 frontmatter 對齊",
    ])
    add_paragraph(doc, "禁止：")
    add_bullets(doc, [
        "把完成的 group 留在 backlog.md「之後再清」",
        "只搬部分子項（要嘛整組完成→整組搬；要嘛部分→留原處）",
    ])

    add_heading(doc, "9. 文件衝突解析順序", 1)
    add_numbered(doc, [
        "當下對話的開發者明確指令",
        "docs/tasks/active.md",
        "docs/memory/current.md",
        "已接受的 ADR（docs/adr/*）",
        "docs/security.md",
        "docs/architecture.md",
        "Session / archive 歷史 log",
    ])

    add_heading(doc, "10. ADR Trigger Checklist", 1)
    add_paragraph(doc, "符合任一項就必須先建立 ADR（status=提議），等開發者標為接受才能進實作：")
    add_bullets(doc, [
        "新增 / 移除核心依賴",
        "Async model / IPC / queue / actor 邊界變更",
        "安全邊界、路徑、網路權限變更",
        "資料格式 / schema / public contract 變更",
        "重大演算法、indexing、儲存模型變更",
    ])

    add_heading(doc, "11. docs:refresh 自動化腳本建議", 1)
    add_paragraph(doc, "建議在 package.json（或對應建置系統）加入：")
    add_code(doc, """{
  "scripts": {
    "docs:sync":    "node scripts/docs-sync.mjs",
    "docs:guard":   "node scripts/docs-guard.mjs",
    "docs:refresh": "npm run docs:sync && npm run docs:guard"
  }
}""")
    add_paragraph(doc, "docs:sync 應做的事：")
    add_bullets(doc, [
        "掃描 docs/**/*.md，從 frontmatter 自動產生 / 更新 docs/index.md 的 Document Map 表格",
        "正規化 CRLF（Windows 易產生 BOM / 重複 frontmatter 問題）",
        "驗證每個 .md 都有合法 frontmatter（type / status / updated 三項至少）",
    ])
    add_paragraph(doc, "docs:guard 應做的事：")
    add_bullets(doc, [
        "在 pre-commit / CI 偵測：本次 commit 改了 src/ 但沒改 docs/tasks/* 或 docs/memory/current.md → 阻擋",
        "偵測 backlog.md 出現 100% [x] 的 group → 阻擋並要求搬到 completed.md",
        "偵測 ADR 被 AI 改成「接受」→ 阻擋",
    ])

    add_heading(doc, "12. 新專案初始化 Checklist", 1)
    add_bullets(doc, [
        "[ ] 建立 CLAUDE.md（根目錄）→ 用第 3 節模板，填入 [PROJECT_NAME] / [STACK] / [BUILD_CMD]",
        "[ ] 建立 docs/CLAUDE.md → 用第 4 節模板",
        "[ ] 建立 docs/index.md → 用第 6.1 節模板",
        "[ ] 建立 docs/project.md → 用第 6.2 節模板",
        "[ ] 建立 docs/memory.md + docs/memory/current.md → 第 6.3 / 6.4 節",
        "[ ] 建立 docs/tasks.md + tasks/{active,backlog,blocked,completed}.md → 第 6.5–6.9 節",
        "[ ] 建立 docs/architecture.md / security.md / testing.md / decisions.md 大綱 → 第 6.10–6.13 節",
        "[ ] 建立 docs/adr/0000-template.md → 第 6.14 節",
        "[ ] 建立第一個實際 ADR（例如 0001 選擇技術棧）",
        "[ ] 設定 docs:refresh 腳本（第 11 節）",
        "[ ] Commit message 必須引用 docs/tasks/* 或 docs/memory/current.md",
        "[ ] 啟動 Claude Code 驗證 session start 能正確讀到 index → current → active",
    ])

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("— End of Template — 修改本模板時請維持 frontmatter 與目錄樹一致 —")
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x90, 0xA5)

    doc.save(OUTPUT)
    print(f"WROTE: {OUTPUT}")


if __name__ == "__main__":
    build_doc()